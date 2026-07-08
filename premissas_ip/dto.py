"""
Extração de premissas a partir do **DTO** (Diagnóstico Técnico Operacional, .docx).

Lê as tabelas do documento e auto-preenche os parâmetros de alta confiança casando
por **palavras-chave estáveis** nos rótulos (não por índice de tabela, que varia entre
municípios) — assim funciona para qualquer DTO no mesmo padrão.

Tudo o que é extraído fica visível para conferência na página; nada é aplicado de forma
silenciosa. Distribuições por classe/tecnologia são apenas detectadas e exibidas (não
aplicadas automaticamente), pois exigem decisão de mapeamento (atual × projetado).

Mantém a seção de Premissas independente: o DTO é um upload local desta página.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field

from . import schema


# ── Normalização e parsing de valores (pt-BR) ─────────────────────────────────
def _norm(s: str) -> str:
    s = unicodedata.normalize("NFD", str(s))
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", s.lower()).strip()


def num_ptbr(txt: str) -> float | None:
    """Converte número em formato pt-BR ('13.854', '1.254,3', 'R$ 80.650,23')."""
    if txt is None:
        return None
    s = re.sub(r"[^0-9.,]", "", str(txt))
    if not s:
        return None
    if "," in s:                       # vírgula = decimal; ponto = milhar
        s = s.replace(".", "").replace(",", ".")
    elif s.count(".") >= 1:
        # só pontos: se o último grupo tem 3 dígitos, é separador de milhar
        if re.fullmatch(r"\d{1,3}(\.\d{3})+", s):
            s = s.replace(".", "")
    try:
        return float(s)
    except ValueError:
        return None


def pct(txt: str) -> float | None:
    """Converte percentual ('9,52%') em fração (0.0952)."""
    v = num_ptbr(txt)
    return None if v is None else v / 100.0


def anos(txt: str) -> float | None:
    """Extrai anos de vida útil ('11 anos', '1,5 e 2 anos' → 1.5, 'Encerrada' → None)."""
    if txt is None or "encerrad" in _norm(txt):
        return None
    m = re.search(r"\d+[.,]?\d*", _norm(txt))      # primeiro número ('1,5' em '1,5 e 2 anos')
    return num_ptbr(m.group(0)) if m else None


# ── Mapa de extração: palavras-chave → parâmetro ──────────────────────────────
@dataclass(frozen=True)
class Regra:
    param_id: str
    contem: tuple[str, ...]          # todas devem estar presentes (normalizadas)
    exclui: tuple[str, ...] = ()     # nenhuma pode estar presente
    conv: str = "num"                # num | pct | anos


_CONV = {"num": num_ptbr, "pct": pct, "anos": anos}

MAPA: tuple[Regra, ...] = (
    Regra("p37", ("total de pontos", "revis"), conv="num"),          # parque total revisado
    Regra("p75", ("led", "revis"), exclui=("total de pontos",), conv="num"),  # LEDs já instalados
    Regra("p349", ("expansao anual",), conv="num"),                  # expansão anual projetada
    Regra("p361", ("demanda reprimida",), conv="num"),               # demanda reprimida total
    Regra("p200", ("luminarias led",), conv="anos"),                 # vida útil luminária LED
)


def mapear_linha(label: str, valor_txt: str) -> tuple[str, float] | None:
    """Casa uma linha (rótulo, valor) com um parâmetro do catálogo, ou None."""
    lab = _norm(label)
    for regra in MAPA:
        if all(k in lab for k in regra.contem) and not any(k in lab for k in regra.exclui):
            v = _CONV[regra.conv](valor_txt)
            if v is not None:
                return regra.param_id, v
    return None


# ── Resultado da extração ─────────────────────────────────────────────────────
@dataclass
class ItemDTO:
    param_id: str
    secao: str
    label_param: str
    label_dto: str
    valor_bruto: str
    valor: float


@dataclass
class Distribuicao:
    titulo: str
    linhas: list[tuple[str, str]]    # (rótulo, valor bruto)


@dataclass
class ExtracaoDTO:
    valores: dict[str, float] = field(default_factory=dict)
    itens: list[ItemDTO] = field(default_factory=list)
    distribuicoes: list[Distribuicao] = field(default_factory=list)
    n_tabelas: int = 0
    avisos: list[str] = field(default_factory=list)


def extrair(arquivo) -> ExtracaoDTO:
    """Extrai premissas de um DTO (.docx). `arquivo` é caminho ou file-like."""
    import docx

    s = schema.carregar()
    doc = docx.Document(arquivo)
    res = ExtracaoDTO(n_tabelas=len(doc.tables))

    for t in doc.tables:
        if not t.rows:
            continue
        cab = _norm(t.rows[0].cells[0].text)
        # Tabelas Parâmetro | Valor (2 colunas) → tentar mapear cada linha.
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            casado = mapear_linha(cells[0], cells[1])
            if casado:
                pid, valor = casado
                if pid in res.valores:
                    continue
                p = s.parametro(pid)
                sec = next((sec.nome for sec in s.secoes if p in sec.parametros), "")
                res.valores[pid] = valor
                res.itens.append(ItemDTO(
                    param_id=pid, secao=sec,
                    label_param=p.label if p else pid,
                    label_dto=cells[0], valor_bruto=cells[1], valor=valor,
                ))
        # Distribuições por classe/tecnologia → apenas detectar para conferência.
        if ("classe de iluminacao" in cab or ("tecnologia" in cab and len(t.columns) <= 3)):
            linhas = [(r.cells[0].text.strip(), r.cells[-1].text.strip())
                      for r in t.rows[1:] if r.cells[0].text.strip()]
            if linhas:
                res.distribuicoes.append(Distribuicao(titulo=t.rows[0].cells[0].text.strip(), linhas=linhas))

    if not res.itens:
        res.avisos.append("Nenhum parâmetro de alta confiança foi reconhecido neste DTO. "
                           "Verifique se o documento segue o padrão de tabelas Parâmetro/Valor.")
    return res


__all__ = ["extrair", "mapear_linha", "num_ptbr", "pct", "anos",
           "ExtracaoDTO", "ItemDTO", "Distribuicao", "MAPA", "Regra"]
